# -*- coding: utf-8 -*-
import os
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

formatted_path = r"C:\Users\iznamu\OneDrive - Panyapiwat Institute of Management\CAI 2nd Year 2025\CAI 2.2 2026\PMO1-03-08-2026\E-CMIS\document\as-is to-be\as-is version\12_กจ14_ระบบบริหารกลางและสนับสนุน_E-CMIS_V1_9_formatted.docx"
template_path = r"C:\Users\iznamu\OneDrive - Panyapiwat Institute of Management\CAI 2nd Year 2025\CAI 2.2 2026\PMO1-03-08-2026\E-CMIS\document\as-is to-be\as-is version\template_PEP.docx"

print("--- Starting Verification of Rebuilt Document ---")

if not os.path.exists(formatted_path):
    print("ERROR: Formatted file does not exist!")
    sys.exit(1)

doc = Document(formatted_path)
temp = Document(template_path)

print(f"File: {os.path.basename(formatted_path)}")
print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")
print(f"Total Sections: {len(doc.sections)}")

print("\n1. Verifying Sections & Margins:")
def check_approx(val1, val2, label):
    diff = abs(val1 - val2)
    if diff < 0.01:
        print(f"  OK: {label} is {val1:.2f} inches")
    else:
        print(f"  FAIL: {label} is {val1:.2f} inches (expected {val2:.2f} inches)")

# Section 0 Portrait
s0 = doc.sections[0]
print("Section 0 (Portrait/Cover):")
print("  different_first_page_header_footer:", s0.different_first_page_header_footer)
check_approx(s0.top_margin.inches, 1.08, "top_margin")
check_approx(s0.bottom_margin.inches, 0.89, "bottom_margin")
check_approx(s0.left_margin.inches, 0.75, "left_margin")
check_approx(s0.right_margin.inches, 0.51, "right_margin")

# Section 2 Landscape
s2 = doc.sections[2]
print("\nSection 2 (Landscape):")
check_approx(s2.top_margin.inches, 0.50, "top_margin")
check_approx(s2.bottom_margin.inches, 0.50, "bottom_margin")
check_approx(s2.left_margin.inches, 0.50, "left_margin")
check_approx(s2.right_margin.inches, 0.50, "right_margin")

print("\n2. Verifying Headers & Footers:")
# Section 1 Header (unlinked portrait)
s1 = doc.sections[1]
print("Section 1 Header Linked:", s1.header.is_linked_to_previous)
print("Section 1 Header Paragraphs:", len(s1.header.paragraphs))
for pi, p in enumerate(s1.header.paragraphs):
    print(f"  HP{pi}: {repr(p.text)}")
print("Section 1 Footer Paragraphs:", len(s1.footer.paragraphs))
for pi, p in enumerate(s1.footer.paragraphs):
    print(f"  FP{pi}: {repr(p.text)}")

# Section 2 Header (unlinked landscape)
print("\nSection 2 Header Linked:", s2.header.is_linked_to_previous)
print("Section 2 Header Paragraphs:", len(s2.header.paragraphs))
for pi, p in enumerate(s2.header.paragraphs):
    print(f"  HP{pi}: {repr(p.text)}")
print("Section 2 Footer Paragraphs:", len(s2.footer.paragraphs))
for pi, p in enumerate(s2.footer.paragraphs):
    print(f"  FP{pi}: {repr(p.text)}")

print("\n3. Verifying Cover Page Text & Fonts:")
cover_ok = True
# Check P3 to P12 fonts (excluding empty paragraphs)
check_indices = [3, 4, 5, 6, 7, 8, 9, 11, 12, 13, 14]
for idx in check_indices:
    p = doc.paragraphs[idx]
    print(f"  P{idx} text: {repr(p.text)}")
    for r in p.runs:
        font_name = r.font.name
        if font_name != 'TH Sarabun New':
            print(f"    FAIL: Run '{r.text}' has font '{font_name}' (expected TH Sarabun New)")
            cover_ok = False
        else:
            print(f"    OK: Run '{r.text[:20]}...' uses TH Sarabun New (size={r.font.size.pt if r.font.size else None})")

if cover_ok:
    print("SUCCESS: Cover page rebuilt successfully and matches template fonts!")
else:
    print("FAIL: Cover page has formatting or font errors!")

print("\n4. Verifying Font Styles (Checking for any TH SarabunPSK remaining):")
old_font_count = 0
new_font_count = 0
other_fonts = set()

# Sample check of first 100 paragraphs in body (starting from P15)
for idx, p in enumerate(doc.paragraphs[15:115]):
    for r in p.runs:
        font_name = r.font.name
        if font_name == 'TH SarabunPSK':
            old_font_count += 1
        elif font_name == 'TH Sarabun New':
            new_font_count += 1
        elif font_name is not None:
            other_fonts.add(font_name)

# Sample check of first 10 tables
for t_idx, t in enumerate(doc.tables[:10]):
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    font_name = r.font.name
                    if font_name == 'TH SarabunPSK':
                        old_font_count += 1
                    elif font_name == 'TH Sarabun New':
                        new_font_count += 1
                    elif font_name is not None:
                        other_fonts.add(font_name)

print(f"  Runs with TH Sarabun New: {new_font_count}")
print(f"  Runs with TH SarabunPSK (should be 0): {old_font_count}")
if other_fonts:
    print(f"  Other fonts found: {list(other_fonts)}")

if old_font_count == 0:
    print("\nSUCCESS: All inspected runs successfully migrated to TH Sarabun New!")
else:
    print("\nFAIL: Inspected runs still contain TH SarabunPSK!")

print("\n5. Verifying Heading sizes:")
h1_sizes = set()
h2_sizes = set()
for p in doc.paragraphs:
    if p.style.name == 'Heading 1':
        for r in p.runs:
            if r.text.strip() and r.font.size:
                h1_sizes.add(r.font.size.pt)
    elif p.style.name == 'Heading 2':
        for r in p.runs:
            if r.text.strip() and r.font.size:
                h2_sizes.add(r.font.size.pt)

print(f"  Heading 1 sizes found: {list(h1_sizes)} (Expected: [22.0])")
print(f"  Heading 2 sizes found: {list(h2_sizes)} (Expected: [18.0])")

print("\n--- Verification Complete ---")
