import copy
import sys
from docx import Document
from docx.oxml import OxmlElement

def main():
    doc_path = r'C:\Users\iznamu\OneDrive\Documents\01_เล่มหลัก_E-CMIS_AS-IS_TO-BE_V1.docx'
    doc = Document(doc_path)

    # 1. Find the paragraph for 2.2
    start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if '2.2 ผังกระบวนงาน AS-IS (BPMN Swimming Lane)' in p.text:
            start_idx = i
            break
            
    if start_idx == -1:
        print("Could not find start paragraph")
        sys.exit(1)

    # 2. Save Table 3 XML (the mapping template)
    table_template = None
    table_idx_to_remove = []
    for i, t in enumerate(doc.tables):
        if len(t.rows) > 0 and 'ลำดับ' in t.rows[0].cells[0].text and 'ขั้นตอน' in t.rows[0].cells[1].text:
            if table_template is None:
                table_template = copy.deepcopy(t._element)
            # Mark table for removal if it's table 3 or later
            if i >= 3:
                table_idx_to_remove.append(t)
        elif i >= 3:
            table_idx_to_remove.append(t)

    if not table_template:
        print("Could not find Table 3 template")
        sys.exit(1)

    # 3. Delete paragraphs from start_idx onwards
    for p in doc.paragraphs[start_idx:]:
        p._element.getparent().remove(p._element)

    # 4. Delete tables from table 3 onwards
    for t in table_idx_to_remove:
        t._element.getparent().remove(t._element)

    # Rebuild Section 2.2 and 2.3
    activities = [
        (4, "ระบบรับเรื่องร้องเรียน"),
        (5, "ระบบกระบวนการดำเนินงานกล่าวหา"),
        (6, "ระบบคุ้มครองพยาน"),
        (7, "ระบบมติคณะกรรมการ ป.ป.ท."),
        (8, "ระบบตรวจสอบประวัติบุคคล"),
        (9, "ระบบหมายจับ"),
        (10, "ระบบกฎหมายในทางคดี"),
    ]

    p_heading2_2 = doc.add_paragraph('2.2 ผังกระบวนงาน AS-IS (BPMN Swimming Lane)')
    try:
        p_heading2_2.style = 'Heading 2'
    except:
        pass

    for act_num, act_name in activities:
        idx = act_num - 3
        p = doc.add_paragraph(f"2.2.{idx} ผังกระบวนงาน AS-IS กิจกรรมที่ {act_num} : {act_name}")
        try:
            p.style = 'Heading 3'
        except:
            pass
        
        p_box = doc.add_paragraph(f"[ เว้นพื้นที่สำหรับแทรกรูปผังกระบวนงาน AS-IS (BPMN Swimming Lane) ของ{act_name} ]")
        p_box.alignment = 1 # Center

    p_heading2_3 = doc.add_paragraph('2.3 รายละเอียดขั้นตอน AS-IS')
    try:
        p_heading2_3.style = 'Heading 2'
    except:
        pass

    for act_num, act_name in activities:
        idx = act_num - 3
        p = doc.add_paragraph(f"2.3.{idx} รายละเอียดขั้นตอน AS-IS กิจกรรมที่ {act_num} : {act_name}")
        try:
            p.style = 'Heading 3'
        except:
            pass
        
        new_tbl = copy.deepcopy(table_template)
        doc._body._element.append(new_tbl)
        
        doc.add_paragraph("")

    doc.save(doc_path)
    print("Successfully regenerated the document.")

if __name__ == '__main__':
    main()
