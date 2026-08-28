# -*- coding: utf-8 -*-
"""
Python script to format 12_กจ14_ระบบบริหารกลางและสนับสนุน_E-CMIS_V1_9.docx
using styling, margins, headers, and footers from the template template_PEP.docx.
Includes cover page rebuilding to match the template cover page layout.
"""
import os
import sys
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import parse_xml

sys.stdout.reconfigure(encoding='utf-8')

# === PATHS ===
source_path = r"C:\Users\iznamu\OneDrive - Panyapiwat Institute of Management\CAI 2nd Year 2025\CAI 2.2 2026\PMO1-03-08-2026\E-CMIS\document\as-is to-be\as-is version\12_กจ14_ระบบบริหารกลางและสนับสนุน_E-CMIS_V1_9.docx"
template_path = r"C:\Users\iznamu\OneDrive - Panyapiwat Institute of Management\CAI 2nd Year 2025\CAI 2.2 2026\PMO1-03-08-2026\E-CMIS\document\as-is to-be\as-is version\template_PEP.docx"
image_path = r"C:\Users\iznamu\.gemini\antigravity\brain\43b1ec64-c6ab-4d79-9e39-6b29d942486f\scratch\template_media\image3.png"
output_path = r"C:\Users\iznamu\OneDrive - Panyapiwat Institute of Management\CAI 2nd Year 2025\CAI 2.2 2026\PMO1-03-08-2026\E-CMIS\document\as-is to-be\as-is version\12_กจ14_ระบบบริหารกลางและสนับสนุน_E-CMIS_V1_9_formatted.docx"

print("Loading documents...")
doc = Document(source_path)
temp = Document(template_path)

# === 1. FONT TRANSLATION HELPER ===
def apply_font_family(run):
    run.font.name = 'TH Sarabun New'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'TH Sarabun New')
    rFonts.set(qn('w:hAnsi'), 'TH Sarabun New')
    rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    rFonts.set(qn('w:cs'), 'TH Sarabun New')

# === 2. REBUILD COVER PAGE ===
print("Rebuilding cover page to match template layout...")
# Cover page paragraphs data (P0 to P14)
cover_data = [
    {"text": "", "size": 16.0, "bold": False, "align": 1, "before": None, "after": None},
    {"text": "", "size": 16.0, "bold": False, "align": 1, "before": None, "after": None},
    {"text": "", "size": 16.0, "bold": False, "align": 1, "before": None, "after": None},
    {"text": "โครงการพัฒนาระบบบริหารจัดการและติดตามสำนวนคดีทุจริต", "size": 26.0, "bold": True, "align": 1, "before": None, "after": None},
    {"text": "และประพฤติมิชอบในภาครัฐ", "size": 26.0, "bold": True, "align": 1, "before": None, "after": None},
    {"text": "Electronic Case Management Intelligence System (E-CMIS)", "size": 22.0, "bold": True, "align": 1, "before": None, "after": 6.0},
    {"text": "กระบวนงานบริหารกลางและสนับสนุน กิจกรรมที่ 14", "size": 26.0, "bold": True, "align": 1, "before": 6.0, "after": None},
    {"text": "ระบบบริหารกลางและสนับสนุน", "size": 20.0, "bold": True, "align": 1, "before": None, "after": 12.0},
    {"text": "สัญญาเลขที่ ๑๔/๒๕๖๙", "size": 22.0, "bold": False, "align": 1, "before": None, "after": 6.0},
    {"text": "โดย", "size": 22.0, "bold": False, "align": 1, "before": None, "after": None},
    {"text": "", "size": 16.0, "bold": False, "align": 1, "before": 6.0, "after": 6.0},
    {"text": "กิจการค้าร่วม SS CONSORTIUM", "size": 22.0, "bold": True, "align": 1, "before": None, "after": None},
    {"text": "(บริษัท สามารถคอมเทค จำกัด / บริษัท สมาร์ทเทอร์แวร์ จำกัด)", "size": 19.0, "bold": False, "align": 1, "before": None, "after": None},
    {"text": "เอกสารนี้จัดทำตามสัญญาเลขที่ ๑๔/๒๕๖๙ เพื่อส่งมอบในงวดที่ 1 ของโครงการ E-CMIS", "size": 14.0, "bold": False, "align": 1, "before": 12.0, "after": None},
    {"text": "ประวัติการปรับปรุงเอกสาร", "size": 22.0, "bold": True, "align": 1, "before": 12.0, "after": 12.0}
]

first_p = doc.paragraphs[0]
for item in cover_data:
    new_p = first_p.insert_paragraph_before()
    new_p.alignment = item["align"]
    pf = new_p.paragraph_format
    if item["before"]:
        pf.space_before = Pt(item["before"])
    if item["after"]:
        pf.space_after = Pt(item["after"])
        
    if item["text"]:
        run = new_p.add_run(item["text"])
        apply_font_family(run)
        run.font.size = Pt(item["size"])
        run.bold = item["bold"]
    else:
        run = new_p.add_run()
        apply_font_family(run)

# Delete old cover paragraphs. Originally doc had P0 to P12 (13 paragraphs) before revision history.
# Now they are at indices 15 to 27.
for idx in range(27, 14, -1):
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)


# === 3. COPY MARGINS ===
print("Applying section margins from template...")
s0_temp = temp.sections[0]
s1_temp = temp.sections[1] # Landscape section in template

# Section 0
s0 = doc.sections[0]
s0.top_margin = s0_temp.top_margin
s0.bottom_margin = s0_temp.bottom_margin
s0.left_margin = s0_temp.left_margin
s0.right_margin = s0_temp.right_margin
s0.header_distance = s0_temp.header_distance
s0.footer_distance = s0_temp.footer_distance

# Section 1
s1 = doc.sections[1]
s1.top_margin = s0_temp.top_margin
s1.bottom_margin = s0_temp.bottom_margin
s1.left_margin = s0_temp.left_margin
s1.right_margin = s0_temp.right_margin
s1.header_distance = s0_temp.header_distance
s1.footer_distance = s0_temp.footer_distance

# Section 2 (Landscape in source)
s2 = doc.sections[2]
s2.top_margin = s1_temp.top_margin
s2.bottom_margin = s1_temp.bottom_margin
s2.left_margin = s1_temp.left_margin
s2.right_margin = s1_temp.right_margin
s2.header_distance = s1_temp.header_distance
s2.footer_distance = s1_temp.footer_distance

# === 4. SETUP HEADERS & FOOTERS ===
print("Setting up headers and footers...")
s0.different_first_page_header_footer = True
s1.header.is_linked_to_previous = True
s1.footer.is_linked_to_previous = True

s2.header.is_linked_to_previous = False
s2.footer.is_linked_to_previous = False

def copy_hdr_ftr_xml(src_hdr_ftr, dst_hdr_ftr, logo_img_path=None):
    dst_hdr_ftr._element.clear()
    rId_map = {}
    if logo_img_path:
        for rId, rel in src_hdr_ftr.part.rels.items():
            if "image" in rel.reltype:
                rId_new, img_obj = dst_hdr_ftr.part.get_or_add_image(logo_img_path)
                rId_map[rId] = rId_new
                
    for child in src_hdr_ftr._element:
        child_copy = copy.deepcopy(child)
        if rId_map:
            xml_str = child_copy.xml
            for old_rId, new_rId in rId_map.items():
                xml_str = xml_str.replace(f'r:embed="{old_rId}"', f'r:embed="{new_rId}"')
                xml_str = xml_str.replace(f'r:id="{old_rId}"', f'r:id="{new_rId}"')
            child_copy = parse_xml(xml_str)
        dst_hdr_ftr._element.append(child_copy)

# Copy portrait header/footer to Section 0
copy_hdr_ftr_xml(temp.sections[0].header, s0.header, image_path)
copy_hdr_ftr_xml(temp.sections[0].footer, s0.footer)

# Copy landscape header/footer to Section 2
copy_hdr_ftr_xml(temp.sections[1].header, s2.header, image_path)
copy_hdr_ftr_xml(temp.sections[1].footer, s2.footer)


# === 5. PARAGRAPH STYLING ===
print("Processing paragraph fonts, sizes, colors and spacing...")
for idx, p in enumerate(doc.paragraphs):
    # Cover page has already been rebuilt and formatted, skip styling
    is_cover_page = idx < 15
    if is_cover_page:
        continue
        
    style_name = p.style.name
    target_size = 16.0  # Default size (User-approved standard for TH Sarabun New)
    target_color = "3C3C3C"  # Default charcoal gray from template
    target_bold = None
    
    # Check if heading
    if style_name == 'Heading 1':
        target_size = 22.0
        target_color = "1B3A6B"  # Dark blue
        target_bold = True
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif style_name == 'Heading 2':
        target_size = 18.0
        target_color = "2E75B6"  # Medium blue
        target_bold = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    elif style_name == 'Heading 3':
        target_size = 16.0
        target_color = "1F4D78"  # Dark blue-gray
        target_bold = True
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
    elif style_name == 'Heading 4':
        target_size = 15.0
        target_color = "2E74B5"
        target_bold = True
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
    else:
        # Normal body text & list paragraphs
        target_size = 16.0
        target_color = "3C3C3C"
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.33

    # Apply properties to runs in this paragraph
    for r in p.runs:
        apply_font_family(r)
        
        # We only override sizes if it's not empty text
        if r.text.strip() or r.text == '\t':
            r.font.size = Pt(target_size)
            if target_color:
                r.font.color.rgb = RGBColor.from_string(target_color)
            if target_bold is not None:
                r.bold = target_bold

# === 6. TABLE STYLING ===
print("Processing tables formatting...")
for t_idx, t in enumerate(doc.tables):
    for r_idx, row in enumerate(t.rows):
        is_header_row = (r_idx == 0)
        for c_idx, cell in enumerate(row.cells):
            is_table0 = (t_idx == 0)
            
            # Format cell paragraphs
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.15
                
                # Apply run styles
                for r in p.runs:
                    apply_font_family(r)
                    
                    if is_table0:
                        r.font.size = Pt(16.0)
                        if c_idx == 0:
                            r.bold = True
                    else:
                        r.font.size = Pt(14.0)
                        if is_header_row:
                            r.bold = True

# === 7. SAVE OUTPUT ===
print(f"Saving formatted document to: {output_path}")
doc.save(output_path)
print("Formatting complete! Cover page successfully rebuilt.")
