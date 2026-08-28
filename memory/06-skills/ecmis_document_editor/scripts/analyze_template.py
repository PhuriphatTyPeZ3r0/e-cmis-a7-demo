# -*- coding: utf-8 -*-
"""Re-analyze the converted template focusing on header/paragraphs with proper encoding."""
import os
from docx import Document
from docx.oxml.ns import qn

base = r'c:\Users\iznamu\OneDrive - Panyapiwat Institute of Management\CAI 2nd Year 2025\CAI 2.2 2026\PMO1-03-08-2026\E-CMIS\document\as-is to-be'

template_dir = None
for d in os.listdir(base):
    full = os.path.join(base, d)
    if os.path.isdir(full) and 'template' in d.lower():
        template_dir = full
        break

converted = os.path.join(template_dir, '_converted_template.docx')
doc = Document(converted)

# Header/Footer analysis
print("=== HEADERS & FOOTERS ===")
for i, section in enumerate(doc.sections):
    header = section.header
    if header and not header.is_linked_to_previous:
        print(f"\nSection {i} HEADER:")
        for pi, p in enumerate(header.paragraphs):
            print(f"  H-P{pi} align={p.alignment}: \"{p.text}\"")
    footer = section.footer
    if footer and not footer.is_linked_to_previous:
        print(f"\nSection {i} FOOTER:")
        for pi, p in enumerate(footer.paragraphs):
            print(f"  F-P{pi} align={p.alignment}: \"{p.text}\"")

# Body paragraphs (full detail)
print("\n=== ALL BODY PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    print(f'P{i} [{p.style.name}] align={p.alignment}: "{p.text}"')
    for j, r in enumerate(p.runs):
        print(f'  R{j}: bold={r.bold}, italic={r.italic}, size={r.font.size}, font={r.font.name}, underline={r.underline}')
        print(f'       text="{r.text}"')

# Raw XML of first few paragraphs to understand structure
print("\n=== RAW XML (first 3 body paragraphs) ===")
for i, p in enumerate(doc.paragraphs[:3]):
    print(f"\nP{i} XML:")
    print(p._element.xml[:500])

# Check if there's a logo/image in header
print("\n=== HEADER IMAGES ===")
for section in doc.sections:
    header = section.header
    for rel in header.part.rels.values():
        if "image" in rel.reltype:
            print(f"  Header image: {rel.target_ref}")
