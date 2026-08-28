# -*- coding: utf-8 -*-
import os
from docx import Document

base = r'c:\Users\iznamu\OneDrive - Panyapiwat Institute of Management\CAI 2nd Year 2025\CAI 2.2 2026\PMO1-03-08-2026\E-CMIS\document\as-is to-be'

template_dir = None
for d in os.listdir(base):
    full = os.path.join(base, d)
    if os.path.isdir(full) and 'template' in d.lower():
        template_dir = full
        break

converted = os.path.join(template_dir, '_converted_template.docx')
doc = Document(converted)

print("=== BODY RUN COLORS ===")
for pi, p in enumerate(doc.paragraphs[:10]):
    for ri, r in enumerate(p.runs):
        color = r.font.color
        color_rgb = color.rgb if color else None
        print(f"P{pi} R{ri} '{r.text}': color={color_rgb}")

print("\n=== TABLE CELL COLORS ===")
if doc.tables:
    t = doc.tables[0]
    for ri in [0, 1, 21, 22]:
        if ri < len(t.rows):
            for ci, cell in enumerate(t.rows[ri].cells):
                for pi, p in enumerate(cell.paragraphs):
                    for run_i, r in enumerate(p.runs):
                        color = r.font.color
                        color_rgb = color.rgb if color else None
                        print(f"Row {ri} Col {ci} P{pi} R{run_i} '{r.text}': color={color_rgb}")
