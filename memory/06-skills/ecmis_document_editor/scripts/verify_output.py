# -*- coding: utf-8 -*-
"""Verify the generated cover pages."""
import os
from docx import Document

base = r'c:\Users\iznamu\OneDrive - Panyapiwat Institute of Management\CAI 2nd Year 2025\CAI 2.2 2026\PMO1-03-08-2026\E-CMIS\document\as-is to-be'

template_dir = None
for d in os.listdir(base):
    full = os.path.join(base, d)
    if os.path.isdir(full) and 'template' in d.lower():
        template_dir = full
        break

# Find output dir
output_dir = None
for d in os.listdir(template_dir):
    full = os.path.join(template_dir, d)
    if os.path.isdir(full) and '13' in d:
        output_dir = full
        break

print(f"Output dir: {output_dir}")
print(f"Files in output:")
for f in sorted(os.listdir(output_dir)):
    print(f"  {f}")

# Verify first and last file
files = sorted([f for f in os.listdir(output_dir) if f.endswith('.docx')])

for idx in [0, 5, 11, 12]:
    if idx < len(files):
        fpath = os.path.join(output_dir, files[idx])
        doc = Document(fpath)
        print(f"\n{'='*60}")
        print(f"FILE: {files[idx]}")
        print(f"{'='*60}")
        for i, p in enumerate(doc.paragraphs[:10]):
            print(f"  P{i}: \"{p.text}\"")
        
        # Table header
        if doc.tables:
            t = doc.tables[0]
            row0 = [t.rows[0].cells[ci].text for ci in range(len(t.rows[0].cells))]
            print(f"  Table header: {row0}")
