import docx, sys
sys.stdout.reconfigure(encoding='utf-8')

f_arm = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้-20260817-update อีกแล้วTT[arm].docx'
f_adj = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้_ปรับปรุงตารางและProcessแล้ว.docx'

d_arm = docx.Document(f_arm)
d_adj = docx.Document(f_adj)

print('=== ARM: Chapter 5 Structure ===')
for i, p in enumerate(d_arm.paragraphs):
    if p.text.strip().startswith('บทที่ 5') or ('5.' in p.text and p.style.name.startswith('Heading')):
        print(f'P{i} [{p.style.name}]: {p.text.strip()}')

print('\n=== ARM: All Tables in Chapter 5 ===')
for idx in range(128, len(d_arm.tables)):
    tbl = d_arm.tables[idx]
    fn = tbl.rows[0].cells[1].text.strip().replace('\n', ' ')
    desc = tbl.rows[5].cells[1].text.strip().replace('\n', ' ')[:80]
    print(f'Tbl {idx}: {fn}\n   Desc: {desc}')
