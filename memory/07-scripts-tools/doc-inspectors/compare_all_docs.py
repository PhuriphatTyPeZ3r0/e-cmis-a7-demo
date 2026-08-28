import docx, sys
sys.stdout.reconfigure(encoding='utf-8')

f_arm = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้-20260817-update อีกแล้วTT[arm].docx'
f_adj = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้_ปรับปรุงตารางและProcessแล้ว.docx'
f_master = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้_รวมตารางMaster_สมบูรณ์.docx'
f_bak = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช-updated-backup-before-tables.docx'

for name, f in [('ARM', f_arm), ('ADJ', f_adj), ('MASTER', f_master), ('BAK', f_bak)]:
    d = docx.Document(f)
    print(f'=== {name}: {f} ===')
    print(f'  Paras: {len(d.paragraphs)}, Tables: {len(d.tables)}')
    print('  Chapter 4 & 5 headings:')
    for i, p in enumerate(d.paragraphs):
        t = p.text.strip()
        if t.startswith('บทที่ 4') or t.startswith('บทที่ 5') or t.startswith('4.') or t.startswith('5.'):
            print(f'    P{i} [{p.style.name}]: {t}')
