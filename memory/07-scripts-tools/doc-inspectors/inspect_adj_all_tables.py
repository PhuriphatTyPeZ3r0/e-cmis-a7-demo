import docx, sys
sys.stdout.reconfigure(encoding='utf-8')

d_adj = docx.Document('เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้_ปรับปรุงตารางและProcessแล้ว.docx')

print('=== D_ADJ CHAPTER 5 TABLES (128 TO END) ===')
for idx in range(128, len(d_adj.tables)):
    tbl = d_adj.tables[idx]
    fn = tbl.rows[0].cells[1].text.strip().replace('\n', ' ')
    menu = tbl.rows[1].cells[1].text.strip().replace('\n', ' ')
    page = tbl.rows[2].cells[1].text.strip().replace('\n', ' ')
    ref = tbl.rows[3].cells[1].text.strip().replace('\n', ' ')
    img_cap = tbl.rows[4].cells[1].text.strip().replace('\n', ' ')
    print(f'Tbl {idx:3d} | FN: {fn[:45]:45s} | Ref: {ref[:35]:35s} | Cap: {img_cap[:35]}')
