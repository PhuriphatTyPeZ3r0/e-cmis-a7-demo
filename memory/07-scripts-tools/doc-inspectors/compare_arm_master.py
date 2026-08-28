import docx, sys
sys.stdout.reconfigure(encoding='utf-8')

f_arm = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้-20260817-update อีกแล้วTT[arm].docx'
f_master = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้_รวมตารางMaster_สมบูรณ์.docx'
f_adj = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้_ปรับปรุงตารางและProcessแล้ว.docx'

d_arm = docx.Document(f_arm)
d_master = docx.Document(f_master)

print('=== ARM vs MASTER comparison ===')
print(f'ARM paras: {len(d_arm.paragraphs)}, tables: {len(d_arm.tables)}')
print(f'MASTER paras: {len(d_master.paragraphs)}, tables: {len(d_master.tables)}')

# Compare paragraph texts
diff_p = []
for i in range(min(len(d_arm.paragraphs), len(d_master.paragraphs))):
    if d_arm.paragraphs[i].text.strip() != d_master.paragraphs[i].text.strip():
        diff_p.append((i, d_arm.paragraphs[i].text.strip(), d_master.paragraphs[i].text.strip()))
print(f'Diff paras count: {len(diff_p)}')
for dp in diff_p[:10]:
    print(f'  P{dp[0]}:\n    ARM: {dp[1][:60]}\n    MAS: {dp[2][:60]}')

# Compare table texts
diff_t = []
for i in range(min(len(d_arm.tables), len(d_master.tables))):
    t1 = ' '.join([c.text.strip() for r in d_arm.tables[i].rows for c in r.cells])
    t2 = ' '.join([c.text.strip() for r in d_master.tables[i].rows for c in r.cells])
    if t1 != t2:
        diff_t.append(i)
print(f'Diff tables count: {len(diff_t)}')
