import os, sys, docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
sys.stdout.reconfigure(encoding='utf-8')

def add_image_to_table_128(doc_path):
    doc = docx.Document(doc_path)
    tbl = doc.tables[128]
    row4_cell = tbl.rows[4].cells[1]
    
    # Check if image already present
    blips = row4_cell._tc.xpath('.//a:blip')
    if not blips:
        p = row4_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture('homepage_master.png', width=Inches(6.0))
        
        # Add caption paragraph
        p_cap = row4_cell.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_cap.add_run('รูปภาพที่ 47 หน้าแรก (ยื่นเรื่องร้องเรียนและติดตามสถานะเรื่องร้องเรียนออนไลน์)')
        run_cap.font.name = 'TH Sarabun PSK'
        run_cap.font.size = Pt(14)
        run_cap.bold = True
        
        doc.save(doc_path)
        print(f'Added image to Table 128 in {doc_path}')
    else:
        print(f'Table 128 already has image in {doc_path}')

add_image_to_table_128('เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้-20260817-update อีกแล้วTT[arm].docx')
