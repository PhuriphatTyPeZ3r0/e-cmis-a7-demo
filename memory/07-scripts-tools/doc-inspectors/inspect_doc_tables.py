import os, sys
import docx
sys.stdout.reconfigure(encoding='utf-8')

doc1_path = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้-20260817-update อีกแล้วTT[arm].docx'
doc2_path = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้_ปรับปรุงตารางและProcessแล้ว.docx'

doc1 = docx.Document(doc1_path)
doc2 = docx.Document(doc2_path)

print(f'Doc 1: {len(doc1.paragraphs)} paragraphs, {len(doc1.tables)} tables')
print(f'Doc 2: {len(doc2.paragraphs)} paragraphs, {len(doc2.tables)} tables')

# Let's inspect the tables at the end of Doc 1 (from table 120 onwards)
for i in range(max(0, len(doc1.tables)-15), len(doc1.tables)):
    tbl = doc1.tables[i]
    r0 = [c.text.strip().replace('\n', ' ') for c in tbl.rows[0].cells]
    r1 = [c.text.strip().replace('\n', ' ') for c in tbl.rows[1].cells] if len(tbl.rows) > 1 else []
    print(f'Doc 1 Table {i} ({len(tbl.rows)} rows x {len(tbl.columns)} cols): {r0[:2]} | {r1[:2]}')

# Let's inspect Doc 2 tables from table 125 onwards
print('\n--- Doc 2 sample tables ---')
for i in range(125, min(140, len(doc2.tables))):
    tbl = doc2.tables[i]
    r0 = [c.text.strip().replace('\n', ' ') for c in tbl.rows[0].cells]
    r1 = [c.text.strip().replace('\n', ' ') for c in tbl.rows[1].cells] if len(tbl.rows) > 1 else []
    print(f'Doc 2 Table {i} ({len(tbl.rows)} rows x {len(tbl.columns)} cols): {r0[:2]} | {r1[:2]}')
