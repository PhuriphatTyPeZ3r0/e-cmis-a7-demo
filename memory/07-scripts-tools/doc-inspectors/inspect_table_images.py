import docx, zipfile, io, sys
from PIL import Image
sys.stdout.reconfigure(encoding='utf-8')

f_arm = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้-20260817-update อีกแล้วTT[arm].docx'
f_adj = 'เล่ม 4 - กจ_4_2.1.3_2.1.4 ออกแบบส่วนการติดต่อกับผู้ใช้_ปรับปรุงตารางและProcessแล้ว.docx'

d_arm = docx.Document(f_arm)
d_adj = docx.Document(f_adj)

z_arm = zipfile.ZipFile(f_arm)
z_adj = zipfile.ZipFile(f_adj)

# Let's inspect images in ARM tables 128-137
print('=== IMAGES IN ARM TABLES 128-137 ===')
for idx in range(128, len(d_arm.tables)):
    tbl = d_arm.tables[idx]
    fn = tbl.rows[0].cells[1].text.strip().replace('\n', ' ')
    rids = tbl._tbl.xpath('.//a:blip/@r:embed')
    if rids:
        # find target
        target = d_arm.part.rels[rids[0]].target_ref
        data = z_arm.read('word/' + target)
        im = Image.open(io.BytesIO(data))
        print(f'Table {idx} [{fn}]: {target}, size={im.size}, bytes={len(data)}')
    else:
        print(f'Table {idx} [{fn}]: NO IMAGE')

print('\n=== TOTAL TABLES IN ADJ ===')
print(f'ADJ has {len(d_adj.tables)} tables.')
