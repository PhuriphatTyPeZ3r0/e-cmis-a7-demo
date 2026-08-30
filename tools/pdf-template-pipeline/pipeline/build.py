"""
build.py — Step 2d.

Input : <out_dir>/schema.json  (+ candidates.json for the source PDF path)
Output: <out_dir>/template.docx     layout-faithful DOCX with {field} placeholders
        <out_dir>/preview.html      A4 .doc-paper fragment (repo ECMIS classes)
        <out_dir>/form.html         Bootstrap left-pane form generated from schema
        <out_dir>/index.html        2-pane demo: form + live preview + DOCX export
        <out_dir>/build-report.md    what matched / what did not

DOCX strategy (user decision): convert the PDF with LibreOffice so the layout is
reproduced, then patch placeholders in by text match. Anything that cannot be
matched is listed in build-report.md for a manual pass.
"""

from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
from html import escape
from pathlib import Path

import docx  # python-docx
import pymupdf

SOFFICE_CANDIDATES = [
    "soffice",
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
]
DOTRUN_RE = re.compile(r"[.\u2026\u00b7\uff0e]{4,}")
REPO_CSS_PREFIX = "../../assets/"   # output-template/<doc>/ -> repo root


# ---------------------------------------------------------------------------
# LibreOffice
# ---------------------------------------------------------------------------
def _soffice() -> str:
    for c in SOFFICE_CANDIDATES:
        if shutil.which(c) or Path(c).exists():
            return c
    raise RuntimeError("LibreOffice (soffice) not found; install it first.")


def pdf_to_docx(pdf_path: Path, out_dir: Path) -> Path:
    soffice = _soffice()
    proc = subprocess.run(
        [soffice, "--headless", "--infilter=writer_pdf_import",
         "--convert-to", "docx", "--outdir", str(out_dir), str(pdf_path)],
        capture_output=True, timeout=180, text=True, encoding="utf-8",
        errors="replace",
    )
    if proc.returncode != 0:
        raise RuntimeError(f"LibreOffice failed ({proc.returncode}):\n"
                           f"{proc.stdout}\n{proc.stderr}")
    produced = out_dir / (pdf_path.stem + ".docx")
    if not produced.exists():
        raise RuntimeError(f"LibreOffice did not produce {produced}")
    target = out_dir / "reference.docx"
    if target.exists():
        target.unlink()
    produced.rename(target)
    return target



# ---------------------------------------------------------------------------
# HTML generation
# ---------------------------------------------------------------------------
PLACEHOLDER_RE = re.compile(r"\{([a-z0-9_]+)\}")


def _pdf_lines(pdf_path: Path) -> list[tuple[str, str]]:
    """(alignment, text) per visual line, read from the PDF in reading order.
    Alignment inferred from the line's horizontal centre vs the text column."""
    doc = pymupdf.open(pdf_path)
    out: list[tuple[str, str]] = []
    for page in doc:
        pw = page.rect.width
        d = page.get_text("dict")
        spans_x = [s["bbox"][0]
                   for b in d["blocks"] if b.get("type") == 0
                   for ln in b.get("lines", []) for s in ln.get("spans", [])]
        left = min(spans_x) if spans_x else 0.0
        right = max((s_x for s_x in
                     [s["bbox"][2]
                      for b in d["blocks"] if b.get("type") == 0
                      for ln in b.get("lines", []) for s in ln.get("spans", [])]),
                    default=pw)
        lines = []
        for b in d["blocks"]:
            if b.get("type") != 0:
                continue
            for ln in b.get("lines", []):
                spans = sorted(ln.get("spans", []), key=lambda s: s["bbox"][0])
                # keep multi-space runs: some blanks are literal space padding
                text = "".join(s["text"] for s in spans).rstrip()
                if not text.strip():
                    lines.append((ln["bbox"][1], "left", ""))
                    continue
                x0 = spans[0]["bbox"][0]
                x1 = spans[-1]["bbox"][2]
                mid = (x0 + x1) / 2
                col_mid = (left + right) / 2
                if x0 - left > 40 and abs(mid - col_mid) < 40:
                    align = "center"
                elif right - x1 < 8 and x0 - left > 80:
                    align = "right"
                else:
                    align = "left"
                lines.append((ln["bbox"][1], align, text))
        for _, align, text in sorted(lines, key=lambda t: t[0]):
            out.append((align, text))
        out.append(("left", ""))  # page break spacer
    doc.close()
    return out


def _apply_fields_to_lines(lines: list[tuple[str, str]],
                           fields: list[dict]) -> tuple[list[tuple[str, str]], set]:
    """Substitute {id} placeholders into the plain PDF lines.

    All matches are computed against the *original* line text (with positions),
    then applied non-overlapping left-to-right, so one field's substitution
    never hides another's anchor context. Returns (lines, matched_ids)."""
    text_lines = [list(t) for t in lines]
    # per line index -> list of (start, end, placeholder, field_id)
    edits: dict[int, list] = {i: [] for i in range(len(text_lines))}
    matched: set = set()

    dot_role = {"kbc_director": "ผอ.กบค", "group_director": "ผอ.กลุ่มงาน",
                "typist": "พิมพ"}

    for f in fields:
        src = f["source"]
        anchor = src.get("anchor", "")
        ph = f["placeholder"]
        fid = f["id"]
        before = src.get("ctx_before", "")
        after = src.get("ctx_after", "")

        if anchor:
            need_ctx = len(anchor) <= 3
            # triage context may contain literal "..." ellipses and stray `;
            # keep only the last / first real word for matching.
            pre = re.sub(r"\s+", " ", re.sub(r"[.`…]+", " ", before)).strip()
            pre = pre.rsplit(" ", 1)[-1] if pre else ""
            post = re.sub(r"\s+", " ", re.sub(r"[.`…]+", " ", after)).strip()
            post = post.split(" ", 1)[0] if post else ""
            placed = False
            for i, row in enumerate(text_lines):
                line = row[1]
                for m in re.finditer(re.escape(anchor), line):
                    free = not any(not (m.end() <= s or m.start() >= e)
                                   for s, e, _, _ in edits[i])
                    if not free:
                        continue
                    if need_ctx:
                        # `before` context is on the same line just left of the
                        # anchor; `after` may be on the next line, so it only
                        # breaks ties, it is not required.
                        if pre and pre not in line[:m.start()]:
                            continue
                    edits[i].append((m.start(), m.end(), ph, fid))
                    matched.add(fid)
                    placed = True
                    break
                if placed:
                    break
            continue

        # --- blank markers -------------------------------------------------
        if fid == "doc_date":
            for i, row in enumerate(text_lines):
                line = row[1]
                if line.lstrip().startswith("วันที่"):
                    gm = re.search(r" {3,}", line)
                    if gm:
                        edits[i].append((gm.start(), gm.end(), " " + ph + " ", fid))
                        matched.add(fid)
                        break
            continue

        if fid == "doc_ref_no" or (before and before.rstrip().endswith("/")):
            key = (before or "ที่").strip().strip("`")
            for i, row in enumerate(text_lines):
                line = row[1].strip()
                if line and (line in key or key.endswith(line) or line.endswith("/")):
                    if "ปป" in line or "ที่" in line:
                        edits[i].append((len(row[1]), len(row[1]), " " + ph, fid))
                        matched.add(fid)
                        break
            continue

        role = dot_role.get(fid.rsplit("_", 1)[0], "")
        want_second = fid.endswith("_date")
        for i, row in enumerate(text_lines):
            line = row[1]
            near = role and (role in line or (i and role in text_lines[i - 1][1]))
            if not near:
                continue
            runs = list(DOTRUN_RE.finditer(line))
            if not runs:
                continue
            pick = runs[1] if (want_second and len(runs) >= 2) else runs[0]
            if any(not (pick.end() <= s or pick.start() >= e)
                   for s, e, _, _ in edits[i]):
                continue
            edits[i].append((pick.start(), pick.end(), ph, fid))
            matched.add(fid)
            break

    for i, row in enumerate(text_lines):
        if not edits[i]:
            continue
        line = row[1]
        for s, e, ph, _ in sorted(edits[i], key=lambda t: -t[0]):
            line = line[:s] + ph + line[e:]
        row[1] = line

    return [(a, t) for a, t in text_lines], matched


def _mergefield_html(text: str, labels: dict) -> str:
    def repl(m):
        fid = m.group(1)
        lbl = labels.get(fid, fid)
        return (f'<span class="mergefield" data-field="{escape(fid)}">'
                f'{escape(lbl)}</span>')
    # escape the non-placeholder parts, then inject spans
    parts = []
    last = 0
    for m in PLACEHOLDER_RE.finditer(text):
        parts.append(escape(text[last:m.start()]))
        parts.append(repl(m))
        last = m.end()
    parts.append(escape(text[last:]))
    return "".join(parts)


# ---------------------------------------------------------------------------
# semantic line roles (for porting into order.html's .doc-paper markup)
# ---------------------------------------------------------------------------
NUM_HEAD_RE = re.compile(r"^\s*[๐-๙0-9]+[.．]")
MEMO_HDR_RE = re.compile(r"^\s*(ที่|วันที่|ส่วนราชการ|เรื่อง|เรียน)\b")
SIGN_LINE_RE = re.compile(r"^\s*\(.*\)\s*$")


def _line_role(text: str, align: str, prev_role: str) -> str:
    """Heuristic map: a PDF text line -> an order.html .doc-paper class role.
    Roles: title | memo-hdr | h | sign | indent
    Overridable per line via structure-<doc>.md."""
    t = text.strip()
    if not t:
        return "gap"
    if t == "บันทึกข้อความ":
        return "title"
    if MEMO_HDR_RE.match(t):
        return "memo-hdr"
    if NUM_HEAD_RE.match(t) and len(t) <= 40:
        return "h"
    if SIGN_LINE_RE.match(t) or t in ("ลงนามแล้ว", "ร้อยเอก"):
        return "sign"
    if align == "center" and prev_role in ("sign", "gap", "title"):
        return "sign"
    return "indent"


def _roles_for(lines, out_dir: Path, doc_name: str) -> list[str]:
    """Guess a role per line, then apply overrides from structure-<doc>.md if
    present. Always (re)writes structure-<doc>.md so a human can correct it."""
    guessed: list[str] = []
    prev = "gap"
    for align, text in lines:
        r = _line_role(text, align, prev)
        guessed.append(r)
        if r != "gap":
            prev = r

    override_path = out_dir / f"structure-{doc_name}.md"
    roles = list(guessed)
    if override_path.exists():
        ov: dict[int, str] = {}
        for ln in override_path.read_text(encoding="utf-8").splitlines():
            m = re.match(
                r"^\s*\|?\s*(\d+)\s*\|\s*(title|memo-hdr|h|sign|indent|gap)\s*\|",
                ln)
            if m:
                ov[int(m.group(1))] = m.group(2)
        for i, r in ov.items():
            if 0 <= i < len(roles):
                roles[i] = r

    md = [f"# Structure — {doc_name}", "",
          "role per line for the order.html port. Edit the `role` column and "
          "re-run build to override. Roles: title | memo-hdr | h | sign | indent | gap",
          "", "| # | role | text |", "|---|---|---|"]
    for i, ((_, text), r) in enumerate(zip(lines, roles)):
        md.append(f"| {i} | {r} | {escape(text.strip()[:80]) or '—'} |")
    override_path.write_text("\n".join(md), encoding="utf-8")
    return roles


ROLE_TO_HTML = {
    "title": ('<div class="doc-title" style="font-size:22px">', "</div>"),
    "memo-hdr": ('<div class="doc-memo-hdr">', "</div>"),
    "h": ('<div class="doc-h" style="margin-top:10px">', "</div>"),
    "sign": ('<div class="doc-sign" style="margin-top:14px">', "</div>"),
    "indent": ('<div class="doc-indent">', "</div>"),
}


def gen_order_fragment(lines, schema: dict, out_dir: Path, doc_name: str,
                       meta: dict) -> None:
    """Emit order-fragment.js: one `ECMIS.OrderMemoDocs["id"] = {...}` block that
    order.html renders into its memo-mode #docTabs pane. bodyHtml keeps literal
    {field} tokens; the order.html adapter swaps them for M() mergefield spans."""
    roles = _roles_for(lines, out_dir, doc_name)

    # merge consecutive lines of the same flowing role into one block, so a
    # wrapped paragraph is a single .doc-indent (matches renderMemoDoc()).
    merged: list[tuple[str, str]] = []  # (role, text)
    for (align, text), role in zip(lines, roles):
        t = text.strip()
        if role == "gap" or not t:
            merged.append(("gap", ""))
            continue
        if (merged and merged[-1][0] == role
                and role in ("indent", "memo-hdr")):
            merged[-1] = (role, merged[-1][1] + " " + t)
        else:
            merged.append((role, t))

    html_parts = []
    for role, text in merged:
        if role == "gap":
            html_parts.append('<div class="doc-gap">&nbsp;</div>')
            continue
        # a single-date field that still trails its literal month/year
        text = re.sub(r"(\{doc_date\})\s*[ก-๙]+\s*[๐-๙0-9]{4}", r"\1", text)
        open_t, close_t = ROLE_TO_HTML.get(role, ROLE_TO_HTML["indent"])
        html_parts.append(f"{open_t}{text}{close_t}")
    body_html = "\n".join(html_parts)

    fields_js = [
        {"id": f["id"], "label": f["label"], "type": f["type"],
         "hint": f.get("hint", ""), "placeholder": _placeholder_for(f)}
        for f in schema["fields"]
    ]
    doc_obj = {
        "id": meta["id"],
        "label": meta["label"],
        "runningTitle": meta.get("runningTitle", schema["doc_name"]),
        "docxTemplate": meta["docxTemplate"],
        "fields": fields_js,
        "prefill": meta.get("prefill", {}),
        "bodyHtml": body_html,
    }
    js = (
        "/* AUTO-GENERATED by tools/pdf-template-pipeline (pipeline/build.py).\n"
        f"   Source: {schema['doc_name']}. Do not hand-edit bodyHtml/fields —\n"
        "   edit the triage / structure files and re-run. `prefill` IS meant to\n"
        "   be hand-filled (schema field id -> ECMIS PrefillSources key). */\n"
        "window.ECMIS = window.ECMIS || {};\n"
        "ECMIS.OrderMemoDocs = ECMIS.OrderMemoDocs || {};\n"
        f"ECMIS.OrderMemoDocs[{json.dumps(meta['id'])}] = "
        + json.dumps(doc_obj, ensure_ascii=False, indent=2) + ";\n"
    )
    (out_dir / "order-fragment.js").write_text(js, encoding="utf-8")
    print(f"[build] order-fragment.js  ({len(fields_js)} fields, "
          f"{len(html_parts)} blocks)")


def _placeholder_for(f: dict) -> str:
    t = f["type"]
    if t == "date":
        return "……………………"
    if t == "number":
        return "…"
    lbl = f.get("label", "")
    return lbl if lbl else "……………"


def gen_preview(lines, schema: dict, out: Path) -> str:
    labels = {f["id"]: (f["label"] or f["id"]) for f in schema["fields"]}
    rows = []
    for align, text in lines:
        if not text:
            rows.append('<p class="doc-gap">&nbsp;</p>')
            continue
        cls = "doc-line" + (f" ta-{align}" if align != "left" else "")
        rows.append(f'<p class="{cls}">{_mergefield_html(text, labels)}</p>')
    body = "\n      ".join(rows)
    out.write_text(PREVIEW_HTML.format(
        doc_name=escape(schema["doc_name"]), css=REPO_CSS_PREFIX, body=body),
        encoding="utf-8")
    return body


def gen_template_docx(lines, schema: dict, out: Path) -> None:
    """A clean, fully-templated DOCX built from the same lines that drive the
    preview. Every {id} is present, so docxtemplater export is 100% reliable.
    Not layout-faithful — reference.docx is the faithful copy."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    amap = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT}
    document = docx.Document()
    style = document.styles["Normal"]
    style.font.name = "TH Sarabun New"
    style.font.size = Pt(16)
    for align, text in lines:
        p = document.add_paragraph()
        p.alignment = amap.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        if text:
            run = p.add_run(text)
            if text.strip() == "ลับ":
                run.bold = True
                run.font.color.rgb = docx.shared.RGBColor(0xDB, 0x26, 0x26)
    document.save(out)


def gen_form(schema: dict, out: Path) -> None:
    out.write_text(FORM_HTML.format(
        doc_name=escape(schema["doc_name"]),
        controls="\n        ".join(_form_control(f) for f in schema["fields"])),
        encoding="utf-8")


def _form_control(f: dict) -> str:
    fid, label, ftype = f["id"], escape(f["label"] or f["id"]), f["type"]
    hint = f.get("hint", "")
    hint_html = f'<div class="form-text">{escape(hint)}</div>' if hint else ""
    if ftype == "textarea":
        ctl = (f'<textarea class="form-control" id="{fid}" name="{fid}" '
               f'rows="3"></textarea>')
    else:
        input_type = {"date": "date", "number": "number"}.get(ftype, "text")
        ctl = (f'<input type="{input_type}" class="form-control" '
               f'id="{fid}" name="{fid}">')
    return (f'<div class="col-md-6">\n'
            f'          <label class="form-label" for="{fid}">{label}</label>\n'
            f'          {ctl}\n          {hint_html}\n        </div>')


def gen_index(schema: dict, preview_body: str, out: Path) -> None:
    out.write_text(INDEX_HTML.format(
        doc_name=escape(schema["doc_name"]),
        css=REPO_CSS_PREFIX,
        preview_body=preview_body,
        schema_json=json.dumps(schema, ensure_ascii=False)),
        encoding="utf-8")


# ---------------------------------------------------------------------------
# HTML templates
# ---------------------------------------------------------------------------
PREVIEW_HTML = """<!doctype html>
<html lang="th">
<head>
<meta charset="utf-8">
<title>Preview — {doc_name}</title>
<link href="https://fonts.googleapis.com/css2?family=Sarabun:wght@400;700&display=swap" rel="stylesheet">
<link href="{css}ecmis-app.css" rel="stylesheet">
<link href="{css}a4-ecmis-workspace.css" rel="stylesheet">
<style>
  body {{ background:#5c6470; margin:0; padding:24px; }}
  .doc-paper .mergefield {{ background:#fff4c2; border-bottom:1px solid #d9b53a; padding:0 2px; }}
  .doc-paper .mergefield.filled {{ background:transparent; border-bottom:none; }}
  .doc-line {{ margin:0 0 2px; font-size:16pt; line-height:1.4; }}
  .doc-line.ta-center {{ text-align:center; }}
  .doc-line.ta-right {{ text-align:right; }}
  .doc-gap {{ margin:0 0 10px; }}
</style>
</head>
<body>
  <div class="doc-paper" id="docPaper">
      {body}
  </div>
</body>
</html>
"""

FORM_HTML = """<!doctype html>
<html lang="th">
<head>
<meta charset="utf-8">
<title>Form — {doc_name}</title>
<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
<link href="https://fonts.googleapis.com/css2?family=Sarabun:wght@400;700&display=swap" rel="stylesheet">
<style> body {{ font-family:'Sarabun',sans-serif; padding:24px; max-width:820px; }} </style>
</head>
<body>
  <h1 class="h4 mb-3">{doc_name}</h1>
  <form id="tplForm" class="row g-3">
        {controls}
  </form>
</body>
</html>
"""

INDEX_HTML = """<!doctype html>
<html lang="th">
<head>
<meta charset="utf-8">
<title>{doc_name}</title>
<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
<link href="https://fonts.googleapis.com/css2?family=Sarabun:wght@400;700&display=swap" rel="stylesheet">
<link href="{css}ecmis-app.css" rel="stylesheet">
<link href="{css}a4-ecmis-workspace.css" rel="stylesheet">
<script src="https://cdn.jsdelivr.net/npm/pizzip@3.1.7/dist/pizzip.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/pizzip@3.1.7/dist/pizzip-utils.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/docxtemplater@3.50.0/build/docxtemplater.js"></script>
<style>
  body {{ font-family:'Sarabun',sans-serif; margin:0; }}
  .wrap {{ display:grid; grid-template-columns:minmax(360px,1fr) minmax(0,1.2fr); gap:0; height:100vh; }}
  .pane-form {{ padding:20px; overflow:auto; border-right:3px solid #caa631; }}
  .pane-doc  {{ background:#5c6470; padding:24px; overflow:auto; }}
  .pane-doc .mergefield {{ background:#fff4c2; border-bottom:1px solid #d9b53a; padding:0 2px; }}
  .pane-doc .mergefield.filled {{ background:transparent; border-bottom:none; }}
  .doc-line {{ margin:0 0 2px; font-size:16pt; line-height:1.4; }}
  .doc-line.ta-center {{ text-align:center; }} .doc-line.ta-right {{ text-align:right; }}
  .doc-gap {{ margin:0 0 10px; }}
</style>
</head>
<body>
<div class="wrap">
  <div class="pane-form">
    <h1 class="h5 mb-3">{doc_name}</h1>
    <form id="tplForm" class="row g-3"></form>
    <button id="btnDocx" class="btn btn-outline-primary btn-sm mt-3">
      <i class="fa-solid fa-file-word"></i> ส่งออก DOCX</button>
  </div>
  <div class="pane-doc"><div class="doc-paper" id="docPaper">
      {preview_body}
  </div></div>
</div>
<script>
const SCHEMA = {schema_json};
const TH_DIGITS = "๐๑๒๓๔๕๖๗๘๙";
const TH_MONTHS = ["มกราคม","กุมภาพันธ์","มีนาคม","เมษายน","พฤษภาคม","มิถุนายน",
  "กรกฎาคม","สิงหาคม","กันยายน","ตุลาคม","พฤศจิกายน","ธันวาคม"];
const toThai = s => String(s).replace(/[0-9]/g, d => TH_DIGITS[d]);
function fmtThaiDate(iso) {{
  if (!iso) return "";
  const [y,m,d] = iso.split("-").map(Number);
  if (!y || !m || !d) return iso;
  return toThai(d) + " " + TH_MONTHS[m-1] + " พ.ศ. " + toThai(y + 543);
}}
function fmtValue(f, raw) {{
  if (!raw) return "";
  if (f.type === "date")   return fmtThaiDate(raw);
  if (f.type === "number") return toThai(raw);
  return raw;
}}

// ---- build the left form from the schema -----------------------------------
const form = document.getElementById("tplForm");
for (const f of SCHEMA.fields) {{
  const col = document.createElement("div");
  col.className = "col-md-6";
  const inputType = f.type === "date" ? "date" : f.type === "number" ? "number" : "text";
  col.innerHTML = `<label class="form-label" for="${{f.id}}">${{f.label || f.id}}</label>` +
    (f.type === "textarea"
      ? `<textarea class="form-control" id="${{f.id}}" rows="3"></textarea>`
      : `<input type="${{inputType}}" class="form-control" id="${{f.id}}">`) +
    (f.hint ? `<div class="form-text">${{f.hint}}</div>` : "");
  form.appendChild(col);
}}

// ---- preview markup is inlined above; just keep it in sync ----------------
const docPaper = document.getElementById("docPaper");
form.addEventListener("input", sync);
sync();

function sync() {{
  for (const f of SCHEMA.fields) {{
    const el = document.getElementById(f.id);
    const val = fmtValue(f, el ? el.value : "");
    docPaper.querySelectorAll(`[data-field="${{f.id}}"]`).forEach(span => {{
      span.textContent = val || (f.label || f.id);
      span.classList.toggle("filled", !!val);
    }});
  }}
}}

// ---- DOCX export via docxtemplater + template.docx ------------------------
document.getElementById("btnDocx").addEventListener("click", () => {{
  PizZipUtils.getBinaryContent("template.docx", (err, content) => {{
    if (err) {{ alert("โหลด template.docx ไม่สำเร็จ: " + err); return; }}
    const zip = new PizZip(content);
    const doc = new window.docxtemplater(zip, {{
      paragraphLoop: true, linebreaks: true,
      delimiters: {{ start: "{{", end: "}}" }},
    }});
    const data = {{}};
    for (const f of SCHEMA.fields) {{
      const el = document.getElementById(f.id);
      data[f.id] = fmtValue(f, el ? el.value : "");
    }}
    doc.render(data);
    const blob = doc.getZip().generate({{ type: "blob",
      mimeType: "application/vnd.openxmlformats-officedocument.wordprocessingml.document" }});
    const a = document.createElement("a");
    a.href = URL.createObjectURL(blob);
    a.download = SCHEMA.doc_name + ".docx";
    a.click();
  }});
}});
</script>
</body>
</html>
"""


# order.html tab metadata, keyed by the leading number of the source doc folder
ORDER_DOC_META = {
    "2": {"id": "notify_zone", "label": "แจ้งเขต",
          "runningTitle": "บันทึกแจ้งรายงานการไต่สวนและวินิจฉัยชี้มูล",
          "docxTemplate": "assets/templates/memo-7x-notify-zone.docx"},
    "5": {"id": "transmit_kbc", "label": "ส่ง กบค.",
          "runningTitle": "บันทึกส่งเอกสารให้กลุ่มงานบริหารติดตามคดี",
          "docxTemplate": "assets/templates/memo-7x-transmit-kbc.docx"},
    "6": {"id": "timebar_report", "label": "ขาดอายุความ",
          "runningTitle": "รายงานคดีขาดอายุความ",
          "docxTemplate": "assets/templates/memo-7x-timebar.docx"},
    "4": {"id": "notify_discipline", "label": "แจ้งโทษวินัย (ร่าง)",
          "runningTitle": "หนังสือแจ้งให้พิจารณาโทษทางวินัย (ร่าง)",
          "docxTemplate": "assets/templates/memo-7x-notify-discipline.docx"},
    "7": {"id": "timebar_secgen", "label": "ขาดอายุความ (เลขาฯ) (ร่าง)",
          "runningTitle": "รายงานคดีขาดอายุความ (ถึงเลขาธิการ) (ร่าง)",
          "docxTemplate": "assets/templates/memo-7x-timebar-secgen.docx"},
    "1": {"id": "submit_inquiry", "label": "เสนอไต่สวน (ร่าง)",
          "runningTitle": "บันทึกเสนอรายงานการไต่สวนเบื้องต้น (ร่าง)",
          "docxTemplate": "assets/templates/memo-7x-submit-inquiry.docx"},
    "3": {"id": "ruling_report", "label": "รายงานวินิจฉัยชี้มูล (ร่าง)",
          "runningTitle": "รายงานการไต่สวนเพื่อวินิจฉัยชี้มูลของคณะกรรมการ ป.ป.ท. (ร่าง)",
          "docxTemplate": "assets/templates/memo-7x-ruling-report.docx"},
}


# ---------------------------------------------------------------------------
def build(out_dir: Path) -> int:
    schema = json.loads((out_dir / "schema.json").read_text(encoding="utf-8"))
    cands = json.loads((out_dir / "candidates.json").read_text(encoding="utf-8"))
    pdf_path = Path(cands["source_pdf"])

    print(f"[build] LibreOffice: {pdf_path.name} -> reference.docx")
    pdf_to_docx(pdf_path, out_dir)

    lines, matched = _apply_fields_to_lines(_pdf_lines(pdf_path), schema["fields"])
    all_ids = [f["id"] for f in schema["fields"]]

    preview_body = gen_preview(lines, schema, out_dir / "preview.html")
    gen_template_docx(lines, schema, out_dir / "template.docx")
    gen_form(schema, out_dir / "form.html")
    gen_index(schema, preview_body, out_dir / "index.html")

    doc_name = schema["doc_name"]
    meta = ORDER_DOC_META.get(doc_name.split(".", 1)[0])
    if meta:
        # keep any hand-filled prefill map from a prior emit
        prev = out_dir / "order-fragment.js"
        if prev.exists():
            pm = re.search(r'"prefill":\s*(\{.*?\}),\n\s*"bodyHtml"',
                           prev.read_text(encoding="utf-8"), re.S)
            if pm:
                try:
                    meta = {**meta, "prefill": json.loads(pm.group(1))}
                except json.JSONDecodeError:
                    pass
        gen_order_fragment(lines, schema, out_dir, doc_name, meta)

    missing = [i for i in all_ids if i not in matched]
    rep = [f"# Build report — {schema['doc_name']}", "",
           f"- {len(matched)}/{len(all_ids)} fields auto-placed into preview.html "
           f"and template.docx",
           "- `reference.docx` = LibreOffice layout-faithful copy (no placeholders)",
           "- `template.docx`  = clean skeleton, every `{id}` present, for "
           "docxtemplater export",
           ""]
    if missing:
        rep += [f"## {len(missing)} field(s) NOT auto-placed — add `{{id}}` by hand",
                "Open template.docx (and mark the spot in preview if needed):", ""]
        rep += [f"- `{i}`" for i in missing]
    else:
        rep.append("All fields placed automatically. ✅")
    rep += ["", "| field | placed |", "|---|:--:|"]
    for fid in all_ids:
        rep.append(f"| `{fid}` | {'✅' if fid in matched else '❌ MANUAL'} |")
    (out_dir / "build-report.md").write_text("\n".join(rep), encoding="utf-8")

    print(f"[build] {len(matched)}/{len(all_ids)} fields placed · see build-report.md")
    for name in ("reference.docx", "template.docx", "preview.html", "form.html",
                 "index.html", "build-report.md"):
        print(f"        {out_dir / name}")
    return 0


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("usage: python -m pipeline.build <out_dir>", file=sys.stderr)
        return 2
    return build(Path(argv[1]))


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
